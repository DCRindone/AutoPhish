from gophish import Gophish
from docx import *
from docx.shared import *
from docx.enum.style import WD_STYLE_TYPE
import urllib3
from pathlib import Path

# Variables and lists for tracking event numbers
targets_opened = []
targets_clicked = []
targets_reported = []
targets_submitted = []

# Lists and dicts for holding prepared report data
campaign_results_summary = []

# Output options
report_format = None
output_word_report = None
output_xlsx_report = None
xlsx_header_bg_color = "2d80b7"
xlsx_header_font_color = "#FFFFFF"

# Gain access to GoPhish
api_key = 'c21434defd36f0158cb2803b2947cfec1bb0dd52296d9e645c41ecbbf4bdc27c'
urllib3.disable_warnings()
api = Gophish(api_key, host='https://127.0.0.1:3333/', verify=False)


def get_report(user_in):
    print("do we get here")
    # Campaign Information
    campaign = api.campaigns.get(campaign_id=user_in)
    api.campaigns.complete(user_in)     # Mark Campaign as complete
    cam_name = campaign.name
    cam_status = campaign.status
    launch_date = campaign.launch_date
    completed_date = campaign.completed_date
    cam_url = campaign.url

    # Collect SMTP information
    smtp = campaign.smtp
    cam_from_address = smtp.from_address
    cam_smtp_host = smtp.host
    cam_smtp_name = smtp.name

    # Collect the template information
    template = campaign.template
    cam_subject_line = template.subject
    cam_template_name = template.name
    cam_template_attachments = template.attachments
    if cam_template_attachments == []:
        cam_template_attachments = "None Used"

    # Collect the landing page information
    page = campaign.page
    cam_page_name = page.name
    cam_redirect_url = page.redirect_url
    if cam_redirect_url == "":
        cam_redirect_url = "Not Used"
    cam_capturing_passwords = page.capture_passwords
    cam_capturing_credentials = page.capture_credentials

    # Create counters for enumeration
    sent_counter = 0
    click_counter = 0
    opened_counter = 0
    reported_counter = 0
    submitted_counter = 0

    # Variables and lists for tracking event numbers
    total_unique_opened = 0
    total_unique_clicked = 0
    total_unique_reported = 0
    total_unique_submitted = 0

    # Run through all events and count each of the four basic events
    for event in campaign.timeline:
        if event.message == "Email Sent":
            sent_counter += 1
        elif event.message == "Email Opened":
            opened_counter += 1
            targets_opened.append(event.email)
        elif event.message == "Clicked Link":
            click_counter += 1
            targets_clicked.append(event.email)
        elif event.message == "Submitted Data":
            submitted_counter += 1
            targets_submitted.append(event.email)
        elif event.message == "Email Reported":
            reported_counter += 1
            targets_reported.append(event.email)

    # Collect User information Per Campaign
    for target in campaign.results:
        temp_dict = {}
        # Add all of the recipient's details and results to the temp dictionary
        temp_dict["email"] = target.email
        temp_dict["fname"] = target.first_name
        temp_dict["lname"] = target.last_name

        # Check if this target was recorded as viewing the email (tracking image)
        if target.email in targets_opened:
            temp_dict["opened"] = True
            total_unique_opened += 1
        else:
            temp_dict["opened"] = False
        # Check if this target clicked the link
        if target.email in targets_clicked:
            temp_dict["clicked"] = True
            total_unique_clicked += 1
        else:
            temp_dict["clicked"] = False
        # Check if this target submitted data
        if target.email in targets_submitted:
            temp_dict["submitted"] = True
            total_unique_submitted += 1
        else:
            temp_dict["submitted"] = False
        # Check if this target reported the email
        if target.email in targets_reported:
            temp_dict["reported"] = True
            total_unique_reported += 1
        else:
            temp_dict["reported"] = False
        # Append the temp dictionary to the event summary list
        campaign_results_summary.append(temp_dict)
    user_total = len(campaign_results_summary)
    write_report(cam_name, launch_date, completed_date, cam_status, user_total, sent_counter, opened_counter,
                 click_counter, submitted_counter, reported_counter, campaign_results_summary, cam_smtp_name,
                 cam_from_address, cam_subject_line, cam_capturing_credentials, cam_capturing_passwords,  cam_url,
                 cam_redirect_url, cam_template_attachments)


def write_report(cam_name, launch_date, completed_date, cam_status, user_total, sent_counter, opened_counter,
                 click_counter, submitted_counter, reported_counter, campaign_results_summary, cam_smtp_name,
                 cam_from_address, cam_subject_line, cam_capturing_credentials, cam_capturing_passwords,
                 cam_url, cam_redirect_url, cam_template_attachments):
    """Assemble and output the Word docx file report."""
    # Create document writer using the template and a style editor
    home = str(Path.home())
    path = r"{}\Desktop\GoPhish\Gophish Extensions\template.docx".format(home)
    d = Document(path)
    styles = d.styles

    # Create a custom styles for table cells
    style = styles.add_style("Cell Text", WD_STYLE_TYPE.CHARACTER)
    cell_text = d.styles["Cell Text"]
    cell_text_font = cell_text.font
    cell_text_font.name = "Calibri"
    cell_text_font.size = Pt(12)
    cell_text_font.bold = True
    cell_text_font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    style = styles.add_style("Cell Text Hit", WD_STYLE_TYPE.CHARACTER)
    cell_text_hit = d.styles["Cell Text Hit"]
    cell_text_hit_font = cell_text_hit.font
    cell_text_hit_font.name = "Calibri"
    cell_text_hit_font.size = Pt(12)
    cell_text_hit_font.bold = True
    cell_text_hit_font.color.rgb = RGBColor(0x00, 0x96, 0x00)

    style = styles.add_style("Cell Text Miss", WD_STYLE_TYPE.CHARACTER)
    cell_text_miss = d.styles["Cell Text Miss"]
    cell_text_miss_font = cell_text_miss.font
    cell_text_miss_font.name = "Calibri"
    cell_text_miss_font.bold = True
    cell_text_miss_font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # Create A launch Date
    day_launched_list = []
    temp_date_launch = launch_date.split("T")
    temp_time_launch = temp_date_launch[1].split(".")
    day_launched_list.append(temp_date_launch[0])
    day_launched_list.append(" at ")
    day_launched_list.append(temp_time_launch[0])
    final_launch_date = "".join(day_launched_list)

    # Create a Completed Date
    day_completed_list = []
    temp_date_comp = completed_date.split("T")
    temp_time_comp = temp_date_comp[1].split(".")
    day_completed_list.append(temp_date_comp[0])
    day_completed_list.append(" at ")
    day_completed_list.append(temp_time_comp[0])
    final_completed_date = "".join(day_completed_list)

    # Write a campaign summary at the top of the report
    d.add_heading("Black Talon Security Phishing Report", 1)
    p = d.add_paragraph()
    run = p.add_run("Campaign Overview For: {}".format(cam_name))
    run.bold = True
    # Runs are basically "runs" of text and must be aligned like we want
    # them aligned in the report -- thus they are pushed left

    # Initial Campaign Overview
    if cam_status == "Completed":
        completed_status = "Completed"
    else:
        completed_status = "Still Active"
    p.add_run("""
  Status: {}
  Launched: {}
  Completed: {}
  Emails Sent: {}
  Emails Opened: {}
  Links Clicked: {}
  Submitted Data: {}
  Emails Reported: {}
    """.format(completed_status, final_launch_date, final_completed_date, sent_counter, opened_counter, click_counter,
               submitted_counter, reported_counter))

# TODO add total section Graph? Numbers?

    # Campaign Details
    p = d.add_paragraph()
    run = p.add_run("Campaign Details:")
    run.bold = True
    p.add_run("""
  From: {}
  Subject: {}
  Phish URL: {}
  Redirect URL: {}
  Attachment(s): {}
  Captured Credentials: {}
  Stored Passwords: {}
    """.format(cam_from_address, cam_subject_line, cam_url, cam_redirect_url, cam_template_attachments,
               cam_capturing_credentials, cam_capturing_passwords))

    d.add_page_break()
    p = d.add_paragraph()
    run = p.add_run("The following table summarizes individual activity:")
    run.bold = True

    # Create a table to hold the event summary results
    table = d.add_table(rows=len(campaign_results_summary) + 1, cols=7, style="GoReport")

    header0 = table.cell(0, 0)
    header0.text = ""
    header0.paragraphs[0].add_run("Last", "Cell Text").bold = True

    header1 = table.cell(0, 1)
    header1.text = ""
    header1.paragraphs[0].add_run("First", "Cell Text").bold = True

    header2 = table.cell(0, 2)
    header2.text = ""
    header2.paragraphs[0].add_run("Email", "Cell Text").bold = True

    header3 = table.cell(0, 3)
    header3.text = ""
    header3.paragraphs[0].add_run("Opened", "Cell Text").bold = True

    header4 = table.cell(0, 4)
    header4.text = ""
    header4.paragraphs[0].add_run("Clicked", "Cell Text").bold = True

    header5 = table.cell(0, 5)
    header5.text = ""
    header5.paragraphs[0].add_run("Data Taken", "Cell Text").bold = True

    header6 = table.cell(0, 6)
    header6.text = ""
    header6.paragraphs[0].add_run("Reported", "Cell Text").bold = True

    # Sort campaign summary by each dict's email entry and then create results table
    counter = 1
    ordered_results = sorted(campaign_results_summary, key=lambda k: k['lname'])
    for target in ordered_results:
        lname_cell = table.cell(counter, 0)
        lname_cell.text = "{}".format(target['lname'])

        fname_cell = table.cell(counter, 1)
        fname_cell.text = "{}".format(target['fname'])

        email_cell = table.cell(counter, 2)
        email_cell.text = "{}".format(target['email'])

        temp_cell = table.cell(counter, 3)
        if target['opened']:
            temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
        else:
            temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

        temp_cell = table.cell(counter, 4)
        if target['clicked']:
            temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
        else:
            temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

        temp_cell = table.cell(counter, 5)
        if target['submitted']:
            temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
        else:
            temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")

        temp_cell = table.cell(counter, 6)
        if target['reported']:
            temp_cell.paragraphs[0].add_run(u'\u2713', "Cell Text Hit")
        else:
            temp_cell.paragraphs[0].add_run(u'\u2718', "Cell Text Miss")
        counter += 1

    # Finalize document and save it as the value of output_word_report
    home = str(Path.home())
    path = r"{}\Desktop\GoPhish\GoPhish Reports\{} Phishing Report ({}).docx".format(home, cam_name, temp_date_launch[0])
    d.save(path)
    # d.save(r"C:\Users\EricWebb\Desktop\GoPhish\GoPhish Reports\{} Phishing Report ({}).docx".format(cam_name, temp_date_launch[0]))


def main():
    user_in = input("Enter campaign ID")
    get_report(user_in)


if __name__ == '__main__':
    main()
